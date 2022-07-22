from datetime import datetime
#################보고서 제작#################보고서 제작#################보고서 제작#################보고서 제작#################

# 보고서 제작 class,def 모음
from turtle import end_fill, width
from docx import Document
from docx.enum.section import WD_ORIENTATION

#Cm,Inches,Pt 단위를 사용하기 위한 모듈
from docx.shared import Cm,Inches,Pt

# 문자 스타일 변경
from docx.enum.style import WD_STYLE_TYPE
from pyparsing import col

# table border style
from docx.table import _Cell
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls

# para 정렬
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

# table 정렬
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENTATION

# font color
from docx.shared import RGBColor

# -*- coding: utf-8 -*-

# filename: add_float_picture.py

'''
Implement floating image based on python-docx.
- Text wrapping style: BEHIND TEXT <wp:anchor behindDoc="1">
- Picture position: top-left corner of PAGE `<wp:positionH relativeFrom="page">`.
Create a docx sample (Layout | Positions | More Layout Options) and explore the
source xml (Open as a zip | word | document.xml) to implement other text wrapping
styles and position modes per `CT_Anchor._anchor_xml()`.
'''

from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from matplotlib import style

# 사진 사이즈 및 위치 변경
# refer to docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
                '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
                '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
                '           %s>\n'
                '  <wp:simplePos x="0" y="0"/>\n'
                '  <wp:positionH relativeFrom="page">\n'
                '    <wp:posOffset>%d</wp:posOffset>\n'
                '  </wp:positionH>\n'
                '  <wp:positionV relativeFrom="page">\n'
                '    <wp:posOffset>%d</wp:posOffset>\n'
                '  </wp:positionV>\n'
                '  <wp:extent cx="914400" cy="914400"/>\n'
                '  <wp:wrapNone/>\n'
                '  <wp:docPr id="666" name="unnamed"/>\n'
                '  <wp:cNvGraphicFramePr>\n'
                '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
                '  </wp:cNvGraphicFramePr>\n'
                '  <a:graphic>\n'
                '    <a:graphicData uri="URI not set"/>\n'
                '  </a:graphic>\n'
                '</wp:anchor>' % (nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y))
        )


# refer to docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `w:anchor` element.
    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)


# 셀 마진 변경
# refer to docx.oxml.__init__.py
register_element_cls('wp:anchor', CT_Anchor)

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    usage:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in ["top", "start", "bottom", "end"]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

# 문단 텍스트 입력
def paragraphText(paragraph,text,fontsize,color,alignment,style):
    text = paragraph.add_run(text)
    text.font.size = Pt(fontsize) 
    font = text.font
    font.color.rgb = RGBColor.from_string(color)

    if alignment == 'CENTER':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'RIGHT':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        pass

    if style == 'bold':
        text.bold = True
    else:
        pass

    return text

def lineSpace(doc,inches,space_before,space_after):
    lineSpace = doc.add_paragraph()
    lineSpace.paragraph_format.line_spacing = Inches(inches)
    lineSpace.paragraph_format.space_before = Pt(space_before)
    lineSpace.paragraph_format.space_after = Pt(space_after)
    return lineSpace

# table 제작
def makeTable(paragraph,row,col,alignment,width,height):

    table = paragraph.add_table(rows=row, cols=col)
    if width == None:
        pass
    else:
        set_col_widths(table,width)

    if height == None:
        pass
    else:
        set_col_height(table,height)    

    if alignment == 'CENTER':
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    else:
        pass
    return table

# table 가로 크기 변경
def set_col_widths(table,widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

# table 세로 크기 변경
def set_col_height(table,heights):
    for idx, row in enumerate(table.rows):
        row.height = heights[idx]

# table 전체 스타일 변경
def titleBorder(
                table,
                top_val,
                top_color,
                top_sz,
                bottom_val,
                bottom_color,
                bottom_sz,
                left_val,
                left_color,
                left_sz,
                right_val,
                right_color,
                right_sz
                ):

    tbl = table._tbl # get xml element in table
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), top_val)
        top.set(qn('w:color'), top_color)
        top.set(qn('w:sz'), top_sz) 

        left = OxmlElement('w:left')
        left.set(qn('w:val'), left_val)
        left.set(qn('w:color'), left_color)
        left.set(qn('w:sz'), left_sz)
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), bottom_val)
        bottom.set(qn('w:color'), bottom_color)
        bottom.set(qn('w:sz'), bottom_sz)

        right = OxmlElement('w:right')
        right.set(qn('w:val'), right_val)
        right.set(qn('w:color'), right_color)
        right.set(qn('w:sz'), right_sz)

        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

# table border 부분 변경
# def set_cell_border(cell: _Cell, **kwargs):
def set_cell_border(table,row,col,**kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    cell = table.rows[row].cells[col]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# 셀 개별 배경색 변환
def cellBackColor(table,row,cell,color):
    #GET CELLS XML ELEMENT
    cell_xml_element = table.rows[row].cells[cell]._tc
    #RETRIEVE THE TABLE CELL PROPERTIES
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    #CREATE SHADING OBJECT
    shade_obj = OxmlElement('w:shd')
    #SET THE SHADING OBJECT
    shade_obj.set(qn('w:fill'), color)
    #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
    table_cell_properties.append(shade_obj)

# 셀 텍스트 삽입
def insertTextCell(table,row,col,text,color,fontSize,fontStyle,vertical_alignment,para_alignment,space_before,space_after):
    cell = table.rows[row].cells[col]
    paragraph = cell.paragraphs[0]
    inputText = paragraph.add_run(text)
    font = inputText.font
    font.color.rgb = RGBColor.from_string(color)

    if fontStyle == 'bold':
        inputText.bold = True
    else:
        pass

    if vertical_alignment == 'CENTER':
        cell.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
    elif vertical_alignment == 'TOP':
        cell.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
    elif vertical_alignment == 'BOTTOM':
        cell.vertical_alignment  = WD_ALIGN_VERTICAL.BOTTOM
    elif vertical_alignment == 'BOTH':
        cell.vertical_alignment  = WD_ALIGN_VERTICAL.BOTH
    else:
        pass

    if para_alignment == 'CENTER':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif para_alignment == 'RIGHT':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif para_alignment == 'LEFT':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        pass
    
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after

    font.size = Pt(fontSize)
    return paragraph

# 셀 병합
def cellMerge(table,stdCellList,cellList):
    stdCell = table.cell(stdCellList[0],stdCellList[1])
    for i in cellList:
        mergeCell = table.cell(i[0],i[1])
        stdCell.merge(mergeCell)

# header 박스 만들기
def makeHeaderBox(No,text):
    headerBoxWidth = [Cm(1.5), Cm(0.5), Cm(17)]
    headerBoxHeight = [Cm(1)]
    headerBox = makeTable(doc,row=1,col=3,alignment='CENTER',width=headerBoxWidth,height=headerBoxHeight)
    set_cell_border(
        headerBox,
        row=0,
        col=2,
        top={"val": "nil"},
        bottom={"val":"single", "sz":"15","color":"#F58D22"},
        start={"val":"nil"},
        end={"val":"nil"}
    )

    insertTextCell(
                    headerBox,
                    row=0,
                    col=0,
                    text=No,
                    color='FFFFFF',
                    fontSize=20,
                    fontStyle='bold',
                    vertical_alignment='CENTER',
                    para_alignment='CENTER',
                    space_after=Pt(0),
                    space_before=Pt(0)
                    )
    insertTextCell(
                    headerBox,
                    row=0,
                    col=2,
                    text=text,
                    color='000000',
                    fontSize=20,
                    fontStyle='bold',
                    vertical_alignment='CENTER',
                    para_alignment='LEFT',
                    space_after=Pt(0),
                    space_before=Pt(0)
                    )

    cellBackColor(headerBox, row=0,cell=0,color="F58D22")

# header 반복 함수
def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row


farmNameReport = 'test'
reportYear = [2017,2018,2019,2020,2021]
reportName = "J - cow"

# 문서 생성
doc = Document()

# 문서 전체 폰트 변경
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 전체 페이지 가로세로 설정
current_section = doc.sections[-1]

# current_section.orientation = WD_ORIENTATION.LANDSCAPE
# current_section.page_width = Cm(29.7)
# current_section.page_height = Cm(21.0)

current_section.top_margin = Cm(1)
current_section.bottom_margin = Cm(1)
current_section.left_margin = Cm(1)
current_section.right_margin = Cm(1)



# page1 표지 만들기

sign = doc.add_paragraph()

# 사진의 크기를 Cm 단위로 설정하여 삽입
add_float_picture(sign,'./Service/signPicture/page.png',width=Cm(21.59), height=Cm(27.94),pos_x=Cm(0), pos_y=Cm(0))

date = datetime.today().strftime("%Y-%m-%d")
page1Date = lineSpace(doc,inches=0.8,space_before=0,space_after=0)
paragraphText(sign,f'작성자 : FarmPlace, 작성일 : {date}',fontsize=12,color='000000',alignment='RIGHT',style='bold')
lineSpace(doc,inches=0.6,space_before=0,space_after=0)
page1Text = lineSpace(doc,inches=0.6,space_before=0,space_after=0)
paragraphText(page1Text,f'< {reportName} 농가현황 및 유전능력평가 > ',fontsize=12,color='000000',alignment='CENTER',style='bold')
lineSpace(doc,inches=1,space_before=0,space_after=0)

paraTitle1 = doc.add_paragraph()
title1 = paragraphText(paraTitle1,text=f'{reportName}',fontsize = 35,color='000000',alignment='CENTER',style='bold')

paraTitle2 = doc.add_paragraph()
title2 = paragraphText(paraTitle2,text='한우 컨설팅 보고서',fontsize = 35,color='000000',alignment='CENTER',style='bold')

paraTitle3 = doc.add_paragraph()
title3 = paragraphText(paraTitle3,text=f'{farmNameReport} 농장',fontsize = 25,color='000000',alignment='CENTER',style='bold')

# page1 표지 만들기 완료

# page2 목차 만들기
doc.add_page_break()

# 목차 만들기
# lineSpace(doc,inches=0.3,space_before=0,space_after=0)

indexBoxWidth = [Cm(19.59)]
indexBoxHeight = [Cm(25.94)]
indexBox = makeTable(doc,row=1,col=1,alignment='CENTER',width=indexBoxWidth,height=indexBoxHeight)

# 목차 테두리 완성
titleBorder(
            indexBox,
            top_val = 'single',
            top_color = '#F58D22',
            top_sz = '25',
            bottom_val = 'single',
            bottom_color = '#F58D22',
            bottom_sz = '25',
            left_val = 'single',
            left_color = '#F58D22',
            left_sz = '25',
            right_val = 'single',
            right_color = '#F58D22',
            right_sz = '25'
            )

# title 박스 생성
indexBoxCell0_0 = indexBox.rows[0].cells[0]

titleBoxWidth = [Cm(17)]
titleBoxHeights = [Cm(0.5),Cm(0.5)]
titleBox = makeTable(indexBoxCell0_0,row=2,col=1,alignment='CENTER',width=titleBoxWidth,height=titleBoxHeights)
# titleBox.style = 'Book Title'
titleBoxCell0_0 = titleBox.rows[0].cells[0]

# title 박스 테두리 변경
set_cell_border(
    titleBox,
    row=0,
    col=0,
    top={"val": "nil"},
    bottom={"val": "single", "sz":"20", "color":"#F58D22"},
    start={"val": "nil"},
    end={"val": "nil"},
)

# title 박스 텍스트 입력
insertTextCell(
                titleBox,
                row=0,
                col=0,
                text=f'{reportName} 컨설팅 보고서',
                color='000000',
                fontSize=13,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='RIGHT',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleBox,
                row=1,
                col=0,
                text='C',
                color='843C0C',
                fontSize=20,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='LEFT',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleBox,
                row=1,
                col=0,
                text='ONTENTS',
                color='F58D22',
                fontSize=20,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='LEFT',
                space_after=Pt(0),
                space_before=Pt(0)
                )

# 목차 리스트 table 생성
titleListWidth = [Cm(2),Cm(15)]
titleListHeights = [Cm(1),Cm(0.1),Cm(1),Cm(0.1),Cm(1),Cm(0.1),Cm(1),Cm(0.1),Cm(1),Cm(0.1),Cm(1),Cm(0.1),Cm(1),Cm(0.1),Cm(1),Cm(0.1)]
titleList = makeTable(indexBoxCell0_0,row=15,col=2,alignment='CENTER',width=titleListWidth,height=titleListHeights)

# 목차 리스트 para style 변경
for i in range(15):
    for j in range(2):
        paragraph = titleList.rows[i].cells[j].paragraphs[0]
        paragraph.style = 'Caption'

# 목차 리스트 table cell 병합
stdCellList1 = [1,0]
cellList1 = [[1,1]]
cellMerge(titleList,stdCellList1,cellList1)

stdCellList2 = [3,0]
cellList2 = [[3,1]]
cellMerge(titleList,stdCellList2,cellList2)

stdCellList3 = [5,0]
cellList3 = [[5,1]]
cellMerge(titleList,stdCellList3,cellList3)

stdCellList4 = [7,0]
cellList4 = [[7,1]]
cellMerge(titleList,stdCellList4,cellList4)

stdCellList5 = [9,0]
cellList5 = [[9,1]]
cellMerge(titleList,stdCellList5,cellList5)

stdCellList6 = [11,0]
cellList6 = [[11,1]]
cellMerge(titleList,stdCellList6,cellList6)

stdCellList7 = [13,0]
cellList7 = [[13,1]]
cellMerge(titleList,stdCellList7,cellList7)

# 목차 리스트 table fontsize 설정
indexFontSize = 20
titleFontSize = 15

# 목차 리스트 table 텍스트 입력
insertTextCell(
                titleList,
                row=0,
                col=0,
                text='01',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=0,
                col=1,
                text=' 농가 정보',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=1,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=2,
                col=0,
                text='02',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=2,
                col=1,
                text=' 도축성적 추세 현황',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=3,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=4,
                col=0,
                text='03',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=4,
                col=1,
                text=' 씨수소(KPN정액) 사용 현황',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=5,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=6,
                col=0,
                text='04',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=6,
                col=1,
                text=' 유전능력 현황',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=7,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=8,
                col=0,
                text='05',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=8,
                col=1,
                text=' 개체별 유전능력 유형분류',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=9,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=10,
                col=0,
                text='06',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=10,
                col=1,
                text=' 선발지수 적용 농가개체순위',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=10,
                col=1,
                text=' (도체중우선)',
                color='000000',
                fontSize=12,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=11,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=12,
                col=0,
                text='07',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=12,
                col=1,
                text=' 선발지수 적용 농가개체순위',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )
                
insertTextCell(
                titleList,
                row=12,
                col=1,
                text=' (근내지방우선)',
                color='000000',
                fontSize=12,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=13,
                col=0,
                text=' ',
                color='000000',
                fontSize=1,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=14,
                col=0,
                text='08',
                color='F58D22',
                fontSize=indexFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                titleList,
                row=14,
                col=1,
                text=f' {reportName} 컨설팅 농가 유전능력 평균 현황',
                color='000000',
                fontSize=titleFontSize,
                fontStyle='bold',
                vertical_alignment='CENTER',
                para_alignment=None,
                space_after=Pt(0),
                space_before=Pt(0)
                )

# 목차 리스트 table 테두리 변경
set_cell_border(
    titleList,
    row=0,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=2,
    col=0,
    top={"val": "nil"},
    bottom={"val":"nil"},
    start={"val":"nil"},
    end={"val":"single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=4,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=6,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=8,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=10,
    col=0,
    top={"val": "nil"},
    bottom={"val":"nil"},
    start={"val":"nil"},
    end={"val":"single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=12,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz":"15"},
)

set_cell_border(
    titleList,
    row=14,
    col=0,
    top={"val": "nil"},
    bottom={"val": "nil"},
    start={"val": "nil"},
    end={"val": "single", "sz":"15"},
)

# note table 생성
noteWidth = [Cm(18)]
noteHeights = [Cm(5)]
noteBox = makeTable(indexBoxCell0_0,row=1,col=1,alignment='CENTER',width=noteWidth,height=noteHeights)

titleBorder(
            noteBox,
            top_val = 'single',
            top_color = '#F58D22',
            top_sz = '3',
            bottom_val = 'single',
            bottom_color = '#F58D22',
            bottom_sz = '3',
            left_val = 'single',
            left_color = '#F58D22',
            left_sz = '3',
            right_val = 'single',
            right_color = '#F58D22',
            right_sz = '3'
            )

noteBoxPara1 = noteBox.rows[0].cells[0].paragraphs[0]
noteBoxPara1.paragraph_format.left_indent = Inches(0.1)           
note1 = paragraphText(
                        noteBoxPara1,
                        text='\n* 보고서에 사용된 사육 및 도축 정보의 경우 농가별 이력제 및 축사로 데이터를 기준으로 작성되었습니다.',
                        fontsize = 10,
                        color='000000',
                        alignment='LEFT',
                        style='bold'
                        )

noteBoxPara2 = noteBox.rows[0].cells[0].add_paragraph()
noteBoxPara2 = noteBox.rows[0].cells[0].paragraphs[1]
noteBoxPara2.paragraph_format.left_indent = Inches(0.1)             
note2 = paragraphText(
                        noteBoxPara2,
                        text='* 도축 정보의 경우 해당 ',
                        fontsize = 10,
                        color='000000',
                        alignment='LEFT',
                        style='bold'
                        )
note2 = paragraphText(
                        noteBoxPara2,
                        text='이력제기관(축협, 한우협회)을 통해 전달받은 정보',
                        fontsize = 10,
                        color='FF0000',
                        alignment='LEFT',
                        style='bold'
                        )
note2 = paragraphText(
                        noteBoxPara2,
                        text='를 사용합니다.',
                        fontsize = 10,
                        color='000000',
                        alignment='LEFT',
                        style='bold'
                        )

noteBoxPara3 = noteBox.rows[0].cells[0].add_paragraph()
noteBoxPara3 = noteBox.rows[0].cells[0].paragraphs[2] 
noteBoxPara3.paragraph_format.left_indent = Inches(0.1)   
note3 = paragraphText(
                        noteBoxPara3,
                        text='* 4page에 ',
                        fontsize = 10,
                        color='000000',
                        alignment='LEFT',
                        style='bold'
                        )
note3 = paragraphText(
                        noteBoxPara3,
                        text='도축 그래프가 없는 번식농가의 경우 ',
                        fontsize = 10,
                        color='FF0000',
                        alignment='LEFT',
                        style='bold'
                        )
note3 = paragraphText(
                        noteBoxPara3,
                        text='자가 출하된 거세우 기록이 크롤링 되지 않았기 때문에\n  그래프를 적용할 수 없습니다.',
                        fontsize = 10,
                        color='000000',
                        alignment='LEFT',
                        style='bold'
                        )

# doc.add_page_break()

# page2 목차 만들기 완료

# page3 농가 현황 및 추세 만들기 완료

makeHeaderBox("01","농가 정보")
lineSpace(doc,inches=0.5,space_before=0,space_after=0)

paraTitle = doc.add_paragraph()
paraTitle.paragraph_format.left_indent = Inches(0.3)
paraTitle.paragraph_format.space_before = Pt(0)
paraTitle.paragraph_format.space_after = Pt(0)  

paragraphText(paraTitle,text='▶',fontsize=13,color='F58D22',alignment='NONE',style='bold')
paragraphText(paraTitle,text=' 농가 정보',fontsize=13,color='000000',alignment='NONE',style='bold')

# # 농가정보 table 생성
farmInfoWidth = [Cm(2.5),Cm(15.5)]
farmInfoHeights = [Cm(1),Cm(1)]
farmInfoBox = makeTable(doc,row=2,col=2,alignment='CENTER',width=farmInfoWidth,height=farmInfoHeights)
farmInfoFontSize = 9
farmInfoFontStyle = 'bold'

setCellVal = "dashSmallGap"

titleBorder(
            farmInfoBox,
            top_val = "single",
            top_color = '#F58D22',
            top_sz = '5',
            bottom_val = "single",
            bottom_color = '#F58D22',
            bottom_sz = '5',
            left_val = "single",
            left_color = '#F58D22',
            left_sz = '5',
            right_val = "single",
            right_color = '#F58D22',
            right_sz = '5'
            )

set_cell_border(
    farmInfoBox,
    row=0,
    col=0,
    bottom={"val": setCellVal, "sz":"5"},
    start={"val": "nil"},
)

set_cell_border(
    farmInfoBox,
    row=0,
    col=1,
    bottom={"val": setCellVal, "sz":"5"},
    start={"val": "single", "sz":"5", "color" : "#F58D22"},
    end={"val": "nil"},
)

set_cell_border(
    farmInfoBox,
    row=1,
    col=0,
    top={"val": setCellVal, "sz":"5"},
    start={"val": "nil"},
)

set_cell_border(
    farmInfoBox,
    row=1,
    col=1,
    top={"val": setCellVal, "sz":"5"},
    start={"val": "single", "sz":"5", "color" : "#F58D22"},
    end={"val": "nil"},
)

cellBackColor(farmInfoBox, row=0, cell=0, color = 'FDECDA')

cellBackColor(farmInfoBox, row=1, cell=0, color = 'FDECDA')

insertTextCell(
                farmInfoBox,
                row=0,
                col=0,
                text='이 름',
                color='000000',
                fontSize=farmInfoFontSize,
                fontStyle=farmInfoFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                farmInfoBox,
                row=1,
                col=0,
                text='위 치',
                color='000000',
                fontSize=farmInfoFontSize,
                fontStyle=farmInfoFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

lineSpace(doc,inches=0.5,space_before=0,space_after=0)

paraTitle = doc.add_paragraph()
paraTitle.paragraph_format.left_indent = Inches(0.3)
paraTitle.paragraph_format.space_before = Pt(0)
paraTitle.paragraph_format.space_after = Pt(0)  

paragraphText(paraTitle,text='▶',fontsize=13,color='F58D22',alignment='NONE',style='bold')
paragraphText(paraTitle,text=' 농가 현황',fontsize=13,color='000000',alignment='NONE',style='bold')
lineSpace(doc,inches=0.01,space_before=0,space_after=0)

# # 목차 리스트 table 생성
farmStatusWidth = [Cm(5),Cm(4),Cm(4),Cm(5)]
# farmStatusHeights = [Cm(1),Cm(1)]
farmStatusBox = makeTable(doc,row=5,col=4,alignment='CENTER',width=farmStatusWidth,height=None)
farmStatusFontSize = 9
farmStatusFontStyle = 'bold'

titleBorder(
            farmStatusBox,
            top_val = 'single',
            top_color = '#F58D22',
            top_sz = '5',
            bottom_val = 'single',
            bottom_color = '#F58D22',
            bottom_sz = '5',
            left_val = 'single',
            left_color = '#F58D22',
            left_sz = '5',
            right_val = 'single',
            right_color = '#F58D22',
            right_sz = '5'
            )

insertTextCell(
                farmStatusBox,
                row=0,
                col=0,
                text='구 분',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

superscript = insertTextCell(
                            farmStatusBox,
                            row=0,
                            col=1,
                            text='도축두수',
                            color='000000',
                            fontSize=farmStatusFontSize,
                            fontStyle=farmStatusFontStyle,
                            vertical_alignment='CENTER',
                            para_alignment='CENTER',
                            space_after=Pt(0),
                            space_before=Pt(0)
                            )

sub_text = superscript.add_run('1:')
sub_text.font.size  = Pt(7)
sub_text.font.superscript  = True

insertTextCell(
                farmStatusBox,
                row=0,
                col=3,
                text='사육두수',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                farmStatusBox,
                row=1,
                col=1,
                text='외부',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                farmStatusBox,
                row=1,
                col=2,
                text='내부',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                farmStatusBox,
                row=2,
                col=0,
                text='수',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                farmStatusBox,
                row=3,
                col=0,
                text='암',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                farmStatusBox,
                row=4,
                col=0,
                text='합계',
                color='000000',
                fontSize=farmStatusFontSize,
                fontStyle=farmStatusFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )


cellBackColor(farmStatusBox, row=0, cell=0, color = 'FDECDA')
cellBackColor(farmStatusBox, row=0, cell=1, color = 'FDECDA')
cellBackColor(farmStatusBox, row=0, cell=3, color = 'FDECDA')
cellBackColor(farmStatusBox, row=1, cell=1, color = 'FDECDA')
cellBackColor(farmStatusBox, row=1, cell=2, color = 'FDECDA')
cellBackColor(farmStatusBox, row=2, cell=0, color = 'FDECDA')
cellBackColor(farmStatusBox, row=3, cell=0, color = 'FDECDA')
cellBackColor(farmStatusBox, row=4, cell=0, color = 'FDECDA')

setCellVal = "dashSmallGap"

set_cell_border(
    farmStatusBox,
    row=0,
    col=0,
    start={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=1,
    col=0,
    start={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=2,
    col=0,
    start={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=3,
    col=0,
    start={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=4,
    col=0,
    start={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=0,
    col=3,
    end={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=1,
    col=3,
    end={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=2,
    col=3,
    end={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=3,
    col=3,
    end={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=4,
    col=3,
    end={"val": "nil"},
)

set_cell_border(
    farmStatusBox,
    row=3,
    col=0,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=3,
    col=1,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=3,
    col=2,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=3,
    col=3,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=2,
    col=0,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=2,
    col=1,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=2,
    col=2,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    farmStatusBox,
    row=2,
    col=3,
    bottom={"val": setCellVal, "sz":"5"}
)

# 농가 현황 table cell 병합
stdCellList1 = [0,1]
cellList1 = [[0,2]]
cellMerge(farmStatusBox,stdCellList1,cellList1)

stdCellList2 = [0,0]
cellList2 = [[1,0]]
cellMerge(farmStatusBox,stdCellList2,cellList2)

stdCellList3 = [0,3]
cellList3 = [[1,3]]
cellMerge(farmStatusBox,stdCellList3,cellList3)


paraTitle = doc.add_paragraph()
paraTitle.paragraph_format.left_indent = Inches(0.3)
paraTitle.paragraph_format.space_before = Pt(0)
paraTitle.paragraph_format.space_after = Pt(0)  

paraSuperscript = paragraphText(paraTitle,text='1:',fontsize=7,color='000000',alignment='NONE',style='None')

paraSuperscript.font.size  = Pt(7)
paraSuperscript.font.superscript  = True

paragraphText(paraTitle,text='농가 도축성적 추세 분석을 위해 추적한 정보로 추정치임',fontsize=7,color='000000',alignment='NONE',style='None')


# 도축성적추세현황 header 생성
lineSpace(doc,inches=0.7,space_before=0,space_after=0)
makeHeaderBox("02","도축성적 추세 현황")
lineSpace(doc,inches=0.3,space_before=0,space_after=0)

paraTitle = doc.add_paragraph()
paraTitle.paragraph_format.left_indent = Inches(0.3)
paraTitle.paragraph_format.space_before = Pt(0)
paraTitle.paragraph_format.space_after = Pt(0)  

paragraphText(paraTitle,text='▶',fontsize=13,color='F58D22',alignment='NONE',style='bold')
paragraphText(paraTitle,text=' 도축개월령 현황표',fontsize=13,color='000000',alignment='NONE',style='bold')


# 도축성적추세현황 table 생성
abattMonthWidth = [Cm(4.5),Cm(2.7),Cm(2.7),Cm(2.7),Cm(2.7),Cm(2.7)]
abattMonthHeight = [Cm(1),Cm(1),Cm(1)]
abattMonthBox = makeTable(doc,row=3,col=6,alignment='CENTER',width=abattMonthWidth,height=abattMonthHeight)
abattMonthFontSize = 9
abattMonthFontStyle = 'bold'

titleBorder(
            abattMonthBox,
            top_val = 'single',
            top_color = '#F58D22',
            top_sz = '5',
            bottom_val = 'single',
            bottom_color = '#F58D22',
            bottom_sz = '5',
            left_val = 'single',
            left_color = '#F58D22',
            left_sz = '5',
            right_val = 'single',
            right_color = '#F58D22',
            right_sz = '5'
            )

insertTextCell(
                abattMonthBox,
                row=0,
                col=0,
                text='연도',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=1,
                col=0,
                text='내부 도축개월령',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=2,
                col=0,
                text='외부 도축개월령',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=0,
                col=1,
                text=f'{str(reportYear[0])}',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=0,
                col=2,
                text=f'{str(reportYear[1])}',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=0,
                col=3,
                text=f'{str(reportYear[2])}',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=0,
                col=4,
                text=f'{str(reportYear[3])}',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

insertTextCell(
                abattMonthBox,
                row=0,
                col=5,
                text=f'{str(reportYear[4])}',
                color='000000',
                fontSize=abattMonthFontSize,
                fontStyle=abattMonthFontStyle,
                vertical_alignment='CENTER',
                para_alignment='CENTER',
                space_after=Pt(0),
                space_before=Pt(0)
                )

cellBackColor(abattMonthBox, row=0, cell=0, color = 'FDECDA')
cellBackColor(abattMonthBox, row=0, cell=1, color = 'FDECDA')
cellBackColor(abattMonthBox, row=0, cell=2, color = 'FDECDA')
cellBackColor(abattMonthBox, row=0, cell=3, color = 'FDECDA')
cellBackColor(abattMonthBox, row=0, cell=4, color = 'FDECDA')
cellBackColor(abattMonthBox, row=0, cell=5, color = 'FDECDA')

setCellVal = "dashSmallGap"
set_cell_border(
    abattMonthBox,
    row=2,
    col=0,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=2,
    col=1,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=2,
    col=2,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=2,
    col=3,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=2,
    col=4,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=2,
    col=5,
    top={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=1,
    col=0,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=1,
    col=1,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=1,
    col=2,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=1,
    col=3,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=1,
    col=4,
    bottom={"val": setCellVal, "sz":"5"}
)

set_cell_border(
    abattMonthBox,
    row=1,
    col=5,
    bottom={"val": setCellVal, "sz":"5"}
)

doc.add_page_break()

doc.save('유전체보고서.docx')